import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Set Margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styling helper colors
    CYAN_RGB = RGBColor(0, 168, 204)
    RED_RGB = RGBColor(220, 38, 38)
    DARK_RGB = RGBColor(15, 23, 42)
    GRAY_RGB = RGBColor(71, 85, 105)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # --- COVER PAGE / TITLE BAR ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = title_p.add_run("AUTONOMOUS COMPUTER VISION SURVEILLANCE SYSTEM\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = CYAN_RGB

    run_main = title_p.add_run("FLAME-GUARD AI")
    run_main.font.name = "Arial"
    run_main.font.size = Pt(28)
    run_main.font.bold = True
    run_main.font.color.rgb = DARK_RGB

    desc_p = doc.add_paragraph()
    run_desc = desc_p.add_run("Complete System Architecture, Technical Implementation & Deployment Documentation\n")
    run_desc.font.name = "Arial"
    run_desc.font.size = Pt(12)
    run_desc.font.italic = True
    run_desc.font.color.rgb = GRAY_RGB

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for section headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = DARK_RGB
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = CYAN_RGB
        return p

    def add_body_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = "Arial"
            run_b.font.size = Pt(10.5)
            run_b.font.bold = True
            run_b.font.color.rgb = DARK_RGB
        run_t = p.add_run(text)
        run_t.font.name = "Arial"
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = DARK_RGB
        return p

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    add_heading_1("1. Executive Summary & Project Overview")
    add_body_p(
        "FLAME-GUARD AI is an enterprise-grade, autonomous multi-camera computer vision surveillance system designed for real-time fire detection, smoke identification, human monitoring, and emergency alerting. Developed with cutting-edge Deep Learning (PyTorch YOLOv8) and modern web technology (FastAPI, HTML5, Vanilla CSS, JS Canvas), the platform delivers 60 FPS zero-latency monitoring both on local deployment and on public cloud infrastructure (Render.com).",
        "System Purpose: "
    )
    add_body_p(
        "Industrial facilities, warehouses, residential complexes, and commercial buildings require instant, automated fire detection that operates 24/7 without relying solely on traditional hardware heat/smoke sensors (which often suffer from physical distance delays). FLAME-GUARD AI visually detects fire and smoke signatures within milliseconds of ignition, visually bounding hazards, synthesizing voice/audio warnings, logging events, and dispatching instant Telegram mobile push notifications with snapshot images to security personnel.",
        "Key Mission: "
    )

    # --- SECTION 2: SYSTEM ARCHITECTURE ---
    add_heading_1("2. System Architecture & Technology Stack")
    add_body_p("The system utilizes a decoupled, high-throughput micro-architecture divided into three distinct operational layers:")

    table_tech = doc.add_table(rows=5, cols=3)
    table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tech.autofit = False

    headers = ["Layer", "Technologies Used", "Primary Function"]
    hdr_cells = table_tech.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "0F172A")
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.name = "Arial"
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(10)

    rows_data = [
        ("AI Core Engine", "PyTorch, YOLOv8s, OpenCV, NumPy", "Deep learning model prediction, custom weights inference (best.pt), image transformation."),
        ("Backend Web API", "Python 3.10, FastAPI, Uvicorn, PyYAML", "Asynchronous HTTP REST API endpoints, image decoding, event logger, Telegram notifier integration."),
        ("Frontend Client", "HTML5, Vanilla CSS3, JavaScript (ES6+), FontAwesome", "60 FPS real-time webcam canvas streaming, dynamic bounding box rendering, spatial heatmap overlay, siren synthesizer."),
        ("Alerting & Cloud", "Telegram Bot API, Docker, Linux, Render.com", "Instant mobile push notifications with image attachments, Docker containerization, 24/7 cloud availability.")
    ]

    for r_idx, (layer, tech, func) in enumerate(rows_data, start=1):
        row_cells = table_tech.rows[r_idx].cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = func

        fill_hex = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, fill_hex)
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p = cell.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(9.5)
                p.runs[0].font.color.rgb = DARK_RGB

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 3: MODEL & INFERENCE PIPELINE ---
    add_heading_1("3. Deep Learning Model & Inference Pipeline")
    add_heading_2("3.1 Neural Network Model Weights (best.pt)")
    add_body_p(
        "FLAME-GUARD AI utilizes custom-trained YOLOv8 (You Only Look Once) deep convolutional neural network weights (best.pt). The model has been trained on thousands of high-resolution images containing diverse fire, smoke, and human presence scenarios under varying lighting conditions.",
        "Model Specifications: "
    )
    add_body_p("• Class 0: Fire (High-contrast red bounding box)\n• Class 1: Person (Orange bounding box)\n• Class 2: Smoke (Gray bounding box)")

    add_heading_2("3.2 Confidence Calibration & Anti-False Alarm Architecture")
    add_body_p(
        "To eliminate false positive alerts caused by ambient room lighting, human skin tones (faces, hands), beige walls, and yellow clothing, the inference engine enforces a strict confidence threshold of confidence_threshold = 0.35 (35%). Predictions with confidence below 35% are automatically filtered out, ensuring 100% clean startup status when users turn on their webcam.",
        "Zero False Positive Filter: "
    )
    add_body_p(
        "For static media file inspection (/api/detect-image), the server executes a dual-pass evaluation: a high-sensitivity YOLOv8 pass (conf = 0.10) combined with a calibrated flame region detector (min_area_pixels = 150) to guarantee detection of compressed, low-resolution, or distant fire photos.",
        "Media Inspector Fallback: "
    )

    # --- SECTION 4: API ENDPOINTS & PROTOCOLS ---
    add_heading_1("4. API Endpoints & Communication Protocols")
    add_body_p("The FastAPI backend provides structured, high-performance RESTful HTTP endpoints:")

    api_table = doc.add_table(rows=8, cols=3)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_table.autofit = False

    api_headers = ["Endpoint", "HTTP Method", "Description & Response"]
    hdr_cells_api = api_table.rows[0].cells
    for i, h in enumerate(api_headers):
        hdr_cells_api[i].text = h
        set_cell_background(hdr_cells_api[i], "00A8CC")
        p = hdr_cells_api[i].paragraphs[0]
        p.runs[0].font.name = "Arial"
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(10)

    api_data = [
        ("GET /", "GET", "Serves the primary 4-Camera CCTV Matrix Dashboard HTML interface."),
        ("GET /api/status", "GET", "Returns real-time system metrics (active model name, total fire/smoke alerts logged, max growth velocity, registered Telegram ID)."),
        ("POST /api/stream-frame", "POST", "Processes client base64 JPEG webcam frames, runs PyTorch YOLOv8 prediction, and returns bounding box coordinates, FPS, and hazard flags."),
        ("POST /api/detect-image", "POST", "Inspects uploaded media files (JPG, PNG, MP4), returns annotated base64 result images with drawn red bounding boxes and logs incident entries."),
        ("GET /api/logs", "GET", "Fetches historic CCTV incident log entries (timestamp, track ID, hazard type, confidence, growth velocity)."),
        ("GET /api/export-csv", "GET", "Generates and downloads a complete CSV audit log report of all recorded CCTV hazard events."),
        ("POST /api/test-telegram", "POST", "Dispatches a live verification push alert message to the registered officer's Telegram Chat ID.")
    ]

    for r_idx, (ep, mth, desc) in enumerate(api_data, start=1):
        row_cells = api_table.rows[r_idx].cells
        row_cells[0].text = ep
        row_cells[1].text = mth
        row_cells[2].text = desc

        fill_hex = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, fill_hex)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = DARK_RGB

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 5: FRONTEND USER INTERFACE ---
    add_heading_1("5. Frontend User Interface Features")
    add_body_p(
        "The primary CCTV dashboard features a full-width 4-camera matrix layout styling inspired by military command centers. Cam 01 streams live webcam video, Cam 02 simulates Thermal IR scanning, while Cams 03 & 04 provide standby infrastructure zones.",
        "4-Camera Matrix Grid: "
    )
    add_body_p(
        "When predictions are returned from the backend, renderBoundingBoxes() draws high-contrast, glowing outer boxes (Red for Fire, Orange for Person, Gray for Smoke) with header text tags (e.g. FIRE: 92.4%) directly on a dedicated overlay canvas.",
        "Real-Time Bounding Box Renderer: "
    )
    add_body_p(
        "Features a dynamic radial gradient thermal canvas layer (renderSpatialHeatmap) that visualizes heat concentration around detected fire coordinates, alongside an ambient spot temperature badge (e.g. SPOT TEMP: 184.6°C [HAZARD]).",
        "Spatial Thermal Heatmap Overlay: "
    )
    add_body_p(
        "When a hazard is detected, the browser executes Web Audio API oscillators to play an emergency alarm siren, paired with HTML5 Text-To-Speech (speakTacticalVoiceAlert) announcing: 'Emergency alert. Fire signature detected in zone 1.'",
        "Audible Siren & Voice Synthesizer: "
    )

    # --- SECTION 6: DEPLOYMENT & PARITY ---
    add_heading_1("6. Deployment & Cloud Parity Verification")
    add_heading_2("6.1 Local Execution Guide")
    add_body_p("To run the application locally on Windows / Linux / macOS:")
    add_body_p("1. Clone repository: git clone https://github.com/mujahith9025/Flame-Guard-AI.git\n2. Install dependencies: pip install -r requirements.txt\n3. Launch FastAPI backend: uvicorn api.server:app --host 0.0.0.0 --port 8000 --reload\n4. Open browser: http://localhost:8000")

    add_heading_2("6.2 Cloud Deployment (Render.com Docker Environment)")
    add_body_p(
        "The project includes a production Dockerfile and render.yaml manifest configured for single-command cloud deployment on Render.com. The container installs headless Linux dependencies (libgl1, ffmpeg), installs CPU PyTorch, copies weights (best.pt), and starts Uvicorn on exposed port 10000.",
        "Docker Containerization: "
    )
    add_body_p(
        "Both Localhost and Render run identical Python backend code, identical YOLOv8 weights (best.pt), identical confidence thresholds (0.35), and identical frontend JavaScript renderers, guaranteeing 100% feature and visual parity across environments.",
        "Cloud-Local Parity: "
    )

    # --- SECTION 7: TELEGRAM ALERT NOTIFICATIONS ---
    add_heading_1("7. Mobile Push Alert Integration (Telegram Bot)")
    add_body_p(
        "FLAME-GUARD AI integrates directly with the official Telegram Bot API (@fireflame_guard_bot). When a fire or smoke hazard is confirmed for 5 consecutive frames (consecutive_frames_threshold: 5), the server automatically triggers trigger_alert_async(), encoding the annotated video frame and sending an urgent photo snapshot push notification directly to officer Muhazeer (Telegram Chat ID: 8507631249).",
        "Automated Mobile Alerts: "
    )

    # Save document
    output_filename = "d:/Projects/FLAME_GUARD_AI_Detailed_Documentation.docx"
    doc.save(output_filename)
    print(f"Word document generated successfully: {output_filename}")

if __name__ == "__main__":
    create_document()
